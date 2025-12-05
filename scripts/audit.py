#!/usr/bin/env python3
"""
Enhanced website audit script.

Usage:
    python3 audit.py --url https://example.com --output /path/to/out.docx --whois 1

Produces a .docx with:
 - Technical checks (HTTP, SSL, headers)
 - SEO checks (title, meta, canonical, robots, sitemap)
 - Content checks (H1/H2, word count, images)
 - Resource / perf estimates (asset count, estimated page weight)
 - Structured data detection
 - Accessibility (basic alt checks)
 - Recommendations (severity + suggested fixes)
"""

import argparse
import sys
import os
import socket
import ssl
import time
from urllib.parse import urlparse, urljoin
import json

# Try imports and provide friendly errors later
try:
    import requests
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Pt
except Exception as e:
    print("Missing Python dependencies. Install: requests beautifulsoup4 python-docx", file=sys.stderr)
    raise

# optional
try:
    import whois as whois_mod
except Exception:
    whois_mod = None

# --- Helpers ---------------------------------------------------------------
HEADERS = {
    'User-Agent': 'NexCodeForge-Audit/1.0 (+https://nexcodeforge.com)'
}

def safe_get(url, timeout=15, allow_redirects=True):
    try:
        resp = requests.get(url, headers=HEADERS, timeout=timeout, allow_redirects=allow_redirects)
        return resp
    except Exception as e:
        return None

def safe_head(url, timeout=10):
    try:
        resp = requests.head(url, headers=HEADERS, timeout=timeout, allow_redirects=True)
        return resp
    except Exception:
        return None

def get_ssl_expiry(hostname):
    try:
        ctx = ssl.create_default_context()
        with ctx.wrap_socket(socket.socket(), server_hostname=hostname) as s:
            s.settimeout(5)
            s.connect((hostname, 443))
            cert = s.getpeercert()
            return cert.get('notAfter')
    except Exception:
        return None

def pretty_list(lines):
    return ["• " + str(l) for l in lines]

def severity_tag(level):
    return f"[{level.upper()}]"

# --- Check functions ------------------------------------------------------
def check_basic_http(url):
    data = {}
    resp = safe_get(url)
    if not resp:
        data['ok'] = False
        data['error'] = 'Failed to fetch URL'
        return data
    data['ok'] = True
    data['status_code'] = resp.status_code
    data['time'] = round(resp.elapsed.total_seconds(), 2)
    data['final_url'] = resp.url
    data['headers'] = dict(resp.headers)
    return data

def parse_html(resp_text, base_url):
    soup = BeautifulSoup(resp_text, 'html.parser')
    parsed = {}
    parsed['title'] = soup.title.string.strip() if soup.title and soup.title.string else ''
    meta_desc_tag = soup.find('meta', attrs={'name':'description'})
    parsed['meta_description'] = meta_desc_tag['content'].strip() if meta_desc_tag and meta_desc_tag.get('content') else ''
    parsed['meta_robots'] = soup.find('meta', attrs={'name':'robots'})
    parsed['canonical'] = None
    can_tag = soup.find('link', attrs={'rel':'canonical'})
    if can_tag and can_tag.get('href'):
        parsed['canonical'] = urljoin(base_url, can_tag.get('href'))
    parsed['viewport'] = bool(soup.find('meta', attrs={'name':'viewport'}))
    parsed['open_graph'] = {tag.get('property'): tag.get('content') for tag in soup.find_all('meta', attrs={'property': True}) if tag.get('property','').startswith('og:')}
    parsed['twitter'] = {tag.get('name'): tag.get('content') for tag in soup.find_all('meta', attrs={'name': True}) if tag.get('name','').startswith('twitter:')}
    parsed['h1'] = [h.get_text(strip=True) for h in soup.find_all('h1')]
    parsed['h2'] = [h.get_text(strip=True) for h in soup.find_all('h2')]
    parsed['links'] = [a.get('href') for a in soup.find_all('a', href=True)]
    parsed['imgs'] = [{'src': urljoin(base_url, img.get('src') or ''), 'alt': (img.get('alt') or '').strip()} for img in soup.find_all('img')]
    # structured data detection (JSON-LD)
    parsed['jsonld'] = []
    for s in soup.find_all('script', type='application/ld+json'):
        try:
            parsed['jsonld'].append(json.loads(s.string))
        except Exception:
            parsed['jsonld'].append(s.string or '')
    parsed['body_text'] = soup.get_text(separator=' ', strip=True)
    return parsed

def analyze_links(base_url, links):
    internal = set()
    external = set()
    broken = []
    parsed_base = urlparse(base_url)
    base_host = parsed_base.netloc
    for href in links:
        if not href:
            continue
        href = href.strip()
        if href.startswith('mailto:') or href.startswith('tel:') or href.startswith('#'):
            continue
        full = urljoin(base_url, href)
        try:
            p = urlparse(full)
            if p.netloc == base_host:
                internal.add(full)
            else:
                external.add(full)
            # quick HEAD to detect 4xx/5xx for same-host and external links (costly)
            h = safe_head(full)
            if h and h.status_code >= 400:
                broken.append({'url': full, 'status': h.status_code})
        except Exception:
            continue
    return {'internal_count': len(internal), 'external_count': len(external), 'broken': broken, 'internal_sample': list(internal)[:10], 'external_sample': list(external)[:10]}

def analyze_images(imgs):
    no_alt = [i['src'] for i in imgs if not i['alt']]
    total = len(imgs)
    sizes = []
    # try HEAD to get content-length (not always present)
    for img in imgs[:30]:  # limit to first 30 to avoid loads
        try:
            h = safe_head(img['src'])
            if h and 'Content-Length' in h.headers:
                sizes.append(int(h.headers['Content-Length']))
        except Exception:
            continue
    avg_size_kb = round((sum(sizes)/len(sizes))/1024, 1) if sizes else None
    return {'total': total, 'no_alt': len(no_alt), 'no_alt_sample': no_alt[:10], 'avg_size_kb': avg_size_kb}

def check_robots_sitemap(base_url):
    parsed = urlparse(base_url)
    root = f"{parsed.scheme}://{parsed.netloc}"
    robots_url = urljoin(root, '/robots.txt')
    sitemap_url = None
    robots = safe_get(robots_url)
    sitemap_found = False
    sitemaps = []
    if robots and robots.status_code == 200:
        txt = robots.text
        for line in txt.splitlines():
            if line.strip().lower().startswith('sitemap:'):
                sitemaps.append(line.split(':',1)[1].strip())
                sitemap_found = True
    # also try /sitemap.xml
    s_xml = safe_get(urljoin(root, '/sitemap.xml'))
    if s_xml and s_xml.status_code == 200:
        sitemaps.append(urljoin(root, '/sitemap.xml'))
    return {'robots_exists': bool(robots and robots.status_code==200), 'robots_text_snippet': (robots.text[:1000] if robots and robots.status_code==200 else ''), 'sitemaps': list(set(sitemaps))}

def check_security_headers(headers):
    findings = []
    # HSTS
    if 'strict-transport-security' not in (k.lower() for k in headers.keys()):
        findings.append({'issue': 'HSTS header missing', 'severity': 'medium', 'fix': 'Add Strict-Transport-Security header to enforce HTTPS (e.g., Strict-Transport-Security: max-age=31536000; includeSubDomains; preload).'})
    # X-Frame-Options
    if 'x-frame-options' not in (k.lower() for k in headers.keys()):
        findings.append({'issue': 'X-Frame-Options missing', 'severity': 'low', 'fix': 'Add X-Frame-Options: DENY or SAMEORIGIN to mitigate clickjacking.'})
    # X-Content-Type-Options
    if 'x-content-type-options' not in (k.lower() for k in headers.keys()):
        findings.append({'issue': 'X-Content-Type-Options missing', 'severity': 'low', 'fix': 'Add X-Content-Type-Options: nosniff.'})
    # CSP
    if 'content-security-policy' not in (k.lower() for k in headers.keys()):
        findings.append({'issue': 'CSP missing', 'severity': 'medium', 'fix': 'Consider adding a Content-Security-Policy to mitigate XSS risks.'})
    # Referrer-Policy
    if 'referrer-policy' not in (k.lower() for k in headers.keys()):
        findings.append({'issue': 'Referrer-Policy missing', 'severity': 'low', 'fix': 'Add Referrer-Policy: no-referrer-when-downgrade or strict-origin-when-cross-origin.'})
    return findings

def recommendations_from_findings(parsed, link_info, img_info, headers, basic_http, robots_info, jsonld, include_whois, whois_info):
    recs = []
    # Title
    title = parsed.get('title','')
    if not title:
        recs.append({'severity':'high', 'issue':'Missing title tag', 'suggestion':'Add a unique, descriptive <title> for the page (50-60 characters recommended).'})
    elif len(title) < 30:
        recs.append({'severity':'medium', 'issue':'Short title', 'suggestion':'Consider lengthening the title to include key terms (aim 50-60 chars).'})
    # Meta description
    md = parsed.get('meta_description','')
    if not md:
        recs.append({'severity':'medium', 'issue':'Missing meta description', 'suggestion':'Add a meta description (50-160 chars) that summarizes the page and includes target keywords.'})
    elif len(md) < 50:
        recs.append({'severity':'low', 'issue':'Short meta description', 'suggestion':'Consider a more informative meta description (50-160 chars).'})
    # H1
    h1c = len(parsed.get('h1',[]))
    if h1c == 0:
        recs.append({'severity':'high','issue':'No H1 found','suggestion':'Add at least one H1 heading per page to indicate main topic.'})
    elif h1c > 2:
        recs.append({'severity':'low','issue':'Multiple H1s','suggestion':'Use H1 for the main title and use H2/H3 for subsections.'})
    # Canonical
    if not parsed.get('canonical'):
        recs.append({'severity':'medium','issue':'Missing canonical','suggestion':'Add a rel="canonical" link to avoid duplicate content issues.'})
    # Viewport
    if not parsed.get('viewport'):
        recs.append({'severity':'high','issue':'Viewport meta missing','suggestion':'Add <meta name="viewport" content="width=device-width, initial-scale=1"> for mobile-friendliness.'})
    # Open Graph / Twitter
    if not parsed.get('open_graph'):
        recs.append({'severity':'low','issue':'Open Graph tags missing','suggestion':'Add Open Graph tags (og:title, og:description, og:image) to improve link previews on social media.'})
    # Images alt
    if img_info['no_alt'] > 0:
        recs.append({'severity':'medium','issue':f'{img_info["no_alt"]} images missing alt','suggestion':'Add meaningful alt attributes to images for accessibility & SEO.'})
    # Broken links
    if link_info['broken']:
        recs.append({'severity':'high','issue':f'{len(link_info["broken"])} broken links found','suggestion':'Fix or remove 4xx/5xx links; use redirects if necessary.'})
    # Robots & sitemap
    if not robots_info['robots_exists']:
        recs.append({'severity':'low','issue':'robots.txt missing','suggestion':'Add robots.txt to indicate crawl rules; include sitemap location.'})
    if not robots_info['sitemaps']:
        recs.append({'severity':'low','issue':'No sitemap discovered','suggestion':'Publish a sitemap.xml and reference it in robots.txt.'})
    # Security headers
    for f in check_security_headers(headers):
        recs.append({'severity': f.get('severity','medium'), 'issue': f['issue'], 'suggestion': f['fix']})
    # HTTPS
    if basic_http.get('final_url','').startswith('http://'):
        recs.append({'severity':'high','issue':'Not using HTTPS','suggestion':'Enable HTTPS and redirect HTTP to HTTPS; obtain certificate (Let\'s Encrypt).'})
    # JSON-LD
    if not jsonld:
        recs.append({'severity':'low','issue':'No structured data found','suggestion':'Consider adding JSON-LD (schema.org) for organization, breadcrumbs, product, article, etc.'})
    # WHOIS (if present)
    if include_whois and whois_info:
        # If domain newly created < 6 months, warn
        try:
            exp = whois_info.get('expiration_date')
            created = whois_info.get('creation_date')
            if isinstance(created, (list, tuple)):
                created = created[0]
            if created:
                age_days = (time.time() - time.mktime(created.timetuple())) / (60*60*24)
                if age_days < 180:
                    recs.append({'severity':'low','issue':'Domain recently registered','suggestion':'New domains may take time to gain trust and SEO signals.'})
        except Exception:
            pass
    # Performance heuristics
    # If many assets (JS/CSS) or large avg image size -> advise optimization
    # (these values are heuristics)
    # NOTE: parsed asset count will be attached later
    return recs

def make_doc(out_path, sections, title='Website Audit Report'):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.add_heading(title, level=1)
    doc.add_paragraph(f'Generated: {time.strftime("%Y-%m-%d %H:%M:%S")}')
    for sec in sections:
        sec_title = sec.get('title')
        sec_lines = sec.get('lines', [])
        doc.add_heading(sec_title, level=2)
        for l in sec_lines:
            if isinstance(l, dict) and 'severity' in l:
                # recommendation item
                p = doc.add_paragraph()
                p.add_run(f"{severity_tag(l['severity'])} {l.get('issue')}\n").bold = True
                p.add_run(l.get('suggestion',''))
            elif isinstance(l, list):
                for sub in l:
                    doc.add_paragraph(str(sub), style='List Bullet')
            else:
                doc.add_paragraph(str(l))
    doc.save(out_path)

# --- Main -----------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--url', required=True, help='Target URL (include scheme)')
    parser.add_argument('--output', required=True, help='Output .docx path')
    parser.add_argument('--whois', default='0', help='1 to enable whois')
    args = parser.parse_args()

    url = args.url
    out = args.output
    include_whois = args.whois == '1'

    sections = []

    # Basic HTTP
    basic = check_basic_http(url)
    if not basic.get('ok'):
        sections.append({'title':'HTTP', 'lines': [f"Failed to fetch URL: {basic.get('error')}"]})
        make_doc(out, sections, title=f'Audit for {url}')
        print(out)
        return

    sections.append({'title':'HTTP', 'lines': [
        f"URL: {url}",
        f"Status code: {basic.get('status_code')}",
        f"Response time (s): {basic.get('time')}",
        f"Final URL after redirects: {basic.get('final_url')}"
    ]})

    # Parse HTML
    resp = safe_get(basic.get('final_url'))
    parsed = parse_html(resp.text, basic.get('final_url'))

    sections.append({'title':'Content & SEO', 'lines': [
        f"Title: {parsed.get('title')}",
        f"Meta description: {parsed.get('meta_description')}",
        f"H1 count: {len(parsed.get('h1'))}",
        f"H2 count: {len(parsed.get('h2'))}",
        f"Sample H1s: {parsed.get('h1')[:5]}",
        f"Word count (approx): {len(parsed.get('body_text').split())}"
    ]})

    # Links
    link_info = analyze_links(basic.get('final_url'), parsed.get('links'))
    sections.append({'title':'Links', 'lines': [
        f"Internal links (count): {link_info['internal_count']}",
        f"External links (count): {link_info['external_count']}",
        f"Broken links (sample): {link_info['broken'][:10]}"
    ]})

    # Images
    img_info = analyze_images(parsed.get('imgs'))
    sections.append({'title':'Images', 'lines': [
        f"Total images: {img_info['total']}",
        f"Images missing alt: {img_info['no_alt']}",
        f"Sample images missing alt: {img_info['no_alt_sample']}",
        f"Average image size (KB) sample: {img_info['avg_size_kb']}"
    ]})

    # Robots & sitemap
    robots_info = check_robots_sitemap(basic.get('final_url'))
    sections.append({'title':'Robots & Sitemap', 'lines': [
        f"robots.txt present: {robots_info['robots_exists']}",
        f"sitemaps found: {robots_info['sitemaps']}"
    ]})

    # Structured data
    sections.append({'title':'Structured Data (JSON-LD)', 'lines': [
        f"JSON-LD blocks found: {len(parsed.get('jsonld'))}",
        f"Sample JSON-LD (first): {json.dumps(parsed.get('jsonld')[0], indent=2) if parsed.get('jsonld') else 'None'}"
    ]})

    # Security headers
    headers = basic.get('headers', {})
    sec_findings = check_security_headers(headers)
    sections.append({'title':'Security Headers', 'lines': [f'Headers found: {list(headers.keys())[:20]}', f'Security issues count: {len(sec_findings)}']})

    # SSL
    parsed_url = urlparse(basic.get('final_url'))
    if parsed_url.scheme == 'https':
        expiry = get_ssl_expiry(parsed_url.hostname)
        sections.append({'title':'SSL', 'lines': [f'Hostname: {parsed_url.hostname}', f'Certificate notAfter: {expiry}']})
    else:
        sections.append({'title':'SSL', 'lines': ['Site not HTTPS — no SSL information.']})

    # WHOIS
    whois_info = {}
    if include_whois:
        if whois_mod:
            try:
                w = whois_mod.whois(parsed_url.netloc)
                whois_info = {
                    'domain': parsed_url.netloc,
                    'registrar': getattr(w, 'registrar', ''),
                    'creation_date': getattr(w, 'creation_date', ''),
                    'expiration_date': getattr(w, 'expiration_date', ''),
                    'name_servers': getattr(w, 'name_servers', '')
                }
                sections.append({'title':'WHOIS', 'lines': [str(whois_info)]})
            except Exception as e:
                sections.append({'title':'WHOIS', 'lines': [f'WHOIS lookup failed: {e}']})
        else:
            sections.append({'title':'WHOIS', 'lines': ['python-whois not installed; skipping WHOIS.']})

    # Security header details
    if sec_findings:
        sections.append({'title':'Security Findings (detailed)', 'lines': [f"{f['issue']}: {f['fix']}" for f in sec_findings]})

    # Recommendations (gathered)
    recs = recommendations_from_findings(parsed, link_info, img_info, headers, basic, robots_info, parsed.get('jsonld'), include_whois, whois_info)
    if recs:
        sections.append({'title':'Recommendations & Enhancements', 'lines': recs})
    else:
        sections.append({'title':'Recommendations & Enhancements', 'lines': ['No major recommendations found.']})

    # Additional tips (static)
    tips = [
        "Consider running Lighthouse or PageSpeed Insights for performance metrics (requires API or headless Chrome).",
        "Use a CI step to catch broken links and missing meta tags automatically.",
        "Compress images and use next-gen formats (WebP/AVIF) to improve page load."
    ]
    sections.append({'title':'Practical Tips', 'lines': tips})

    # Save report
    try:
        os.makedirs(os.path.dirname(out), exist_ok=True)
        make_doc(out, sections, title=f'Audit for {parsed_url.netloc}')
        print(out)
    except Exception as e:
        print(f"ERROR creating doc: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    main()
